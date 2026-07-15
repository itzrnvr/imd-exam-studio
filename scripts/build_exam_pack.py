from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.fonts import addMapping
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(r"C:\Users\babys\OneDrive\文件\imd")
CONTENT_DIR = ROOT / "content"
PDF_DIR = ROOT / "output" / "pdf"
DOCX_DIR = ROOT / "output" / "doc"

DOCUMENTS = [
    ("imd_study_guide.md", "IMD_6-7_Hour_Study_Guide"),
    ("imd_numerical_workbook.md", "IMD_Numerical_Practice_Workbook"),
    ("imd_mock_papers.md", "IMD_100_Mark_Mock_Papers_With_Solutions"),
    ("imd_formula_sheet.md", "IMD_Last_Hour_Formula_And_Symbol_Sheet"),
]

NAVY = colors.HexColor("#183153")
BLUE = colors.HexColor("#2563EB")
LIGHT_BLUE = colors.HexColor("#EAF2FF")
PALE = colors.HexColor("#F7F9FC")
GOLD = colors.HexColor("#D97706")
TEXT = colors.HexColor("#172033")
MUTED = colors.HexColor("#5C667A")


def register_fonts() -> None:
    candidates = {
        "Arial": Path(r"C:\Windows\Fonts\arial.ttf"),
        "Arial-Bold": Path(r"C:\Windows\Fonts\arialbd.ttf"),
        "Arial-Italic": Path(r"C:\Windows\Fonts\ariali.ttf"),
        "Consolas": Path(r"C:\Windows\Fonts\consola.ttf"),
        "Consolas-Bold": Path(r"C:\Windows\Fonts\consolab.ttf"),
    }
    for name, path in candidates.items():
        if path.exists():
            pdfmetrics.registerFont(TTFont(name, str(path)))
    addMapping("Arial", 0, 0, "Arial")
    addMapping("Arial", 1, 0, "Arial-Bold")
    addMapping("Arial", 0, 1, "Arial-Italic")


def inline_markup(text: str) -> str:
    escaped = html.escape(text, quote=False)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(
        r"`(.+?)`",
        r'<font name="Consolas" color="#173A5E">\1</font>',
        escaped,
    )
    return escaped


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)
    return rows


def split_blocks(markdown_text: str) -> list[tuple[str, object]]:
    lines = markdown_text.replace("\r\n", "\n").split("\n")
    blocks: list[tuple[str, object]] = []
    paragraph: list[str] = []
    index = 0

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            blocks.append(("paragraph", " ".join(s.strip() for s in paragraph)))
            paragraph = []

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped == "```":
            flush_paragraph()
            index += 1
            code: list[str] = []
            while index < len(lines) and lines[index].strip() != "```":
                code.append(lines[index])
                index += 1
            blocks.append(("formula", "\n".join(code)))
            index += 1
            continue

        if stripped == "<!-- PAGEBREAK -->":
            flush_paragraph()
            blocks.append(("pagebreak", ""))
            index += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            blocks.append(("heading", (len(heading.group(1)), heading.group(2))))
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                index += 1
            blocks.append(("table", parse_table(table_lines)))
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            callout: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                callout.append(lines[index].strip()[1:].strip())
                index += 1
            blocks.append(("callout", " ".join(callout)))
            continue

        bullet = re.match(r"^(\s*)[-*]\s+(.+)$", raw)
        if bullet:
            flush_paragraph()
            blocks.append(("bullet", (len(bullet.group(1)), bullet.group(2))))
            index += 1
            continue

        numbered = re.match(r"^(\s*)(\d+)\.\s+(.+)$", raw)
        if numbered:
            flush_paragraph()
            blocks.append(
                (
                    "numbered",
                    (len(numbered.group(1)), numbered.group(2), numbered.group(3)),
                )
            )
            index += 1
            continue

        if stripped in {"---", "***"}:
            flush_paragraph()
            blocks.append(("rule", ""))
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            blocks.append(("spacer", ""))
            index += 1
            continue

        paragraph.append(raw)
        index += 1

    flush_paragraph()
    return blocks


def pdf_styles() -> dict[str, ParagraphStyle]:
    samples = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "TitleCustom",
            parent=samples["Title"],
            fontName="Arial-Bold",
            fontSize=27,
            leading=32,
            textColor=NAVY,
            alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "h1": ParagraphStyle(
            "H1Custom",
            parent=samples["Heading1"],
            fontName="Arial-Bold",
            fontSize=18,
            leading=22,
            textColor=NAVY,
            spaceBefore=12,
            spaceAfter=8,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "H2Custom",
            parent=samples["Heading2"],
            fontName="Arial-Bold",
            fontSize=13,
            leading=16,
            textColor=BLUE,
            spaceBefore=9,
            spaceAfter=5,
            keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "H3Custom",
            parent=samples["Heading3"],
            fontName="Arial-Bold",
            fontSize=10.5,
            leading=13,
            textColor=GOLD,
            spaceBefore=7,
            spaceAfter=3,
            keepWithNext=True,
        ),
        "h4": ParagraphStyle(
            "H4Custom",
            parent=samples["Heading4"],
            fontName="Arial-Bold",
            fontSize=9.4,
            leading=12,
            textColor=NAVY,
            spaceBefore=5,
            spaceAfter=2,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "BodyCustom",
            parent=samples["BodyText"],
            fontName="Arial",
            fontSize=9.2,
            leading=12.2,
            textColor=TEXT,
            alignment=TA_LEFT,
            spaceAfter=4,
        ),
        "bullet": ParagraphStyle(
            "BulletCustom",
            parent=samples["BodyText"],
            fontName="Arial",
            fontSize=9,
            leading=11.5,
            leftIndent=15,
            firstLineIndent=-8,
            textColor=TEXT,
            spaceAfter=2,
        ),
        "numbered": ParagraphStyle(
            "NumberedCustom",
            parent=samples["BodyText"],
            fontName="Arial",
            fontSize=9,
            leading=11.5,
            leftIndent=17,
            firstLineIndent=-12,
            textColor=TEXT,
            spaceAfter=2,
        ),
        "callout": ParagraphStyle(
            "CalloutCustom",
            parent=samples["BodyText"],
            fontName="Arial-Bold",
            fontSize=9,
            leading=12,
            textColor=NAVY,
        ),
        "formula": ParagraphStyle(
            "FormulaCustom",
            parent=samples["BodyText"],
            fontName="Consolas",
            fontSize=8.6,
            leading=11.2,
            textColor=colors.HexColor("#0F3557"),
        ),
        "table": ParagraphStyle(
            "TableCustom",
            parent=samples["BodyText"],
            fontName="Arial",
            fontSize=7.4,
            leading=9.2,
            textColor=TEXT,
        ),
        "table_header": ParagraphStyle(
            "TableHeaderCustom",
            parent=samples["BodyText"],
            fontName="Arial-Bold",
            fontSize=7.5,
            leading=9.4,
            textColor=colors.white,
        ),
    }


class ExamDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, doc_title: str) -> None:
        super().__init__(
            filename,
            pagesize=A4,
            rightMargin=0.62 * inch,
            leftMargin=0.62 * inch,
            topMargin=0.68 * inch,
            bottomMargin=0.62 * inch,
            title=doc_title,
            author="OpenAI Codex - prepared from supplied IMD course resources",
        )
        self.doc_title = doc_title
        frame = Frame(
            self.leftMargin,
            self.bottomMargin,
            self.width,
            self.height,
            id="normal",
        )
        self.addPageTemplates(
            [PageTemplate(id="main", frames=[frame], onPage=self.draw_header_footer)]
        )

    def draw_header_footer(self, canvas, doc) -> None:  # type: ignore[no-untyped-def]
        canvas.saveState()
        width, height = A4
        canvas.setStrokeColor(colors.HexColor("#D5DCE8"))
        canvas.setLineWidth(0.6)
        canvas.line(self.leftMargin, height - 0.43 * inch, width - self.rightMargin, height - 0.43 * inch)
        canvas.setFont("Arial", 7.2)
        canvas.setFillColor(MUTED)
        header = self.doc_title.replace("_", " ")
        if len(header) > 78:
            header = header[:75] + "..."
        canvas.drawString(self.leftMargin, height - 0.32 * inch, header)
        canvas.drawRightString(width - self.rightMargin, 0.34 * inch, f"Page {doc.page}")
        canvas.restoreState()


def build_pdf(markdown_path: Path, output_path: Path, title: str) -> None:
    styles = pdf_styles()
    blocks = split_blocks(markdown_path.read_text(encoding="utf-8"))
    story: list[object] = []
    first_heading = True

    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload  # type: ignore[misc]
            style_key = "title" if first_heading and level == 1 else f"h{level}"
            first_heading = False
            story.append(Paragraph(inline_markup(text), styles[style_key]))
        elif kind == "paragraph":
            story.append(Paragraph(inline_markup(str(payload)), styles["body"]))
        elif kind == "bullet":
            indent, text = payload  # type: ignore[misc]
            style = styles["bullet"].clone(f"Bullet{indent}")
            style.leftIndent = 15 + min(indent, 8) * 3
            story.append(Paragraph(f"&#8226; {inline_markup(text)}", style))
        elif kind == "numbered":
            indent, number, text = payload  # type: ignore[misc]
            style = styles["numbered"].clone(f"Numbered{indent}")
            style.leftIndent = 17 + min(indent, 8) * 3
            story.append(Paragraph(f"{number}. {inline_markup(text)}", style))
        elif kind == "formula":
            formula = html.escape(str(payload), quote=False).replace("\n", "<br/>")
            box = Table(
                [[Paragraph(formula, styles["formula"])]],
                colWidths=[6.9 * inch],
                hAlign="LEFT",
            )
            box.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F0F6FC")),
                        ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#91B4D5")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 9),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.extend([box, Spacer(1, 5)])
        elif kind == "callout":
            callout = Table(
                [[Paragraph(inline_markup(str(payload)), styles["callout"])]],
                colWidths=[6.9 * inch],
                hAlign="LEFT",
            )
            callout.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), LIGHT_BLUE),
                        ("LINEBEFORE", (0, 0), (0, -1), 4, BLUE),
                        ("LEFTPADDING", (0, 0), (-1, -1), 10),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 7),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                    ]
                )
            )
            story.extend([callout, Spacer(1, 5)])
        elif kind == "table":
            rows = payload  # type: ignore[assignment]
            if not rows:
                continue
            column_count = max(len(row) for row in rows)
            normalized = [row + [""] * (column_count - len(row)) for row in rows]
            rendered_rows: list[list[Paragraph]] = []
            for row_index, row in enumerate(normalized):
                style = styles["table_header"] if row_index == 0 else styles["table"]
                rendered_rows.append([Paragraph(inline_markup(cell), style) for cell in row])
            table = Table(
                rendered_rows,
                colWidths=[6.9 * inch / column_count] * column_count,
                repeatRows=1,
                hAlign="LEFT",
            )
            commands = [
                ("BACKGROUND", (0, 0), (-1, 0), NAVY),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#AEB8C7")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
            for row_index in range(1, len(rendered_rows)):
                if row_index % 2 == 0:
                    commands.append(("BACKGROUND", (0, row_index), (-1, row_index), PALE))
            table.setStyle(TableStyle(commands))
            story.extend([table, Spacer(1, 6)])
        elif kind == "pagebreak":
            story.append(PageBreak())
        elif kind == "rule":
            rule = Table([[""]], colWidths=[6.9 * inch], rowHeights=[1])
            rule.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#CDD6E3"))]))
            story.extend([Spacer(1, 2), rule, Spacer(1, 4)])
        elif kind == "spacer":
            story.append(Spacer(1, 2))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = ExamDocTemplate(str(output_path), title)
    document.build(story)


def set_cell_shading(cell, fill: str) -> None:  # type: ignore[no-untyped-def]
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph) -> None:  # type: ignore[no-untyped-def]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_docx(document: Document, title: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor(23, 32, 51)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for index, size, color in [
        (1, 18, RGBColor(24, 49, 83)),
        (2, 13, RGBColor(37, 99, 235)),
        (3, 10.5, RGBColor(217, 119, 6)),
        (4, 9.5, RGBColor(24, 49, 83)),
    ]:
        style = document.styles[f"Heading {index}"]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    if "Formula Block" not in document.styles:
        style = document.styles.add_style("Formula Block", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(8.6)
        style.font.color.rgb = RGBColor(15, 53, 87)
        style.paragraph_format.left_indent = Inches(0.08)
        style.paragraph_format.right_indent = Inches(0.08)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = title.replace("_", " ")
    header.style = document.styles["Normal"]
    for run in header.runs:
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(92, 102, 122)
    add_page_number(section.footer.paragraphs[0])

    document.core_properties.title = title.replace("_", " ")
    document.core_properties.author = "OpenAI Codex"
    document.core_properties.subject = "Intelligent Model Design exam preparation"


def add_inline_runs(paragraph, text: str) -> None:  # type: ignore[no-untyped-def]
    token_pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(23, 58, 94)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def build_docx(markdown_path: Path, output_path: Path, title: str) -> None:
    blocks = split_blocks(markdown_path.read_text(encoding="utf-8"))
    document = Document()
    configure_docx(document, title)
    first_heading = True

    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload  # type: ignore[misc]
            if first_heading and level == 1:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(80)
                paragraph.paragraph_format.space_after = Pt(18)
                run = paragraph.add_run(text)
                run.font.name = "Arial"
                run.font.bold = True
                run.font.size = Pt(26)
                run.font.color.rgb = RGBColor(24, 49, 83)
                first_heading = False
            else:
                document.add_heading(text, level=level)
        elif kind == "paragraph":
            add_inline_runs(document.add_paragraph(), str(payload))
        elif kind == "bullet":
            indent, text = payload  # type: ignore[misc]
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.18 + min(indent, 8) * 0.04)
            add_inline_runs(paragraph, text)
        elif kind == "numbered":
            indent, number, text = payload  # type: ignore[misc]
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.2 + min(indent, 8) * 0.04)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            add_inline_runs(paragraph, f"{number}. {text}")
        elif kind == "formula":
            paragraph = document.add_paragraph(style="Formula Block")
            paragraph.paragraph_format.keep_together = True
            paragraph.add_run(str(payload))
            p_pr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F0F6FC")
            p_pr.append(shd)
        elif kind == "callout":
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            set_cell_shading(cell, "EAF2FF")
            paragraph = cell.paragraphs[0]
            add_inline_runs(paragraph, str(payload))
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(24, 49, 83)
        elif kind == "table":
            rows = payload  # type: ignore[assignment]
            if not rows:
                continue
            column_count = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=column_count)
            table.style = "Table Grid"
            table.autofit = True
            for row_index, row in enumerate(rows):
                for column_index in range(column_count):
                    cell = table.cell(row_index, column_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    text = row[column_index] if column_index < len(row) else ""
                    paragraph = cell.paragraphs[0]
                    add_inline_runs(paragraph, text)
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(7.4)
                    if row_index == 0:
                        set_cell_shading(cell, "183153")
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    elif row_index % 2 == 0:
                        set_cell_shading(cell, "F7F9FC")
        elif kind == "pagebreak":
            document.add_page_break()
        elif kind == "rule":
            paragraph = document.add_paragraph()
            p_pr = paragraph._p.get_or_add_pPr()
            bottom = OxmlElement("w:pBdr")
            border = OxmlElement("w:bottom")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "6")
            border.set(qn("w:color"), "CDD6E3")
            bottom.append(border)
            p_pr.append(bottom)
        elif kind == "spacer":
            continue

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    register_fonts()
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_DIR.mkdir(parents=True, exist_ok=True)

    for source_name, output_stem in DOCUMENTS:
        source_path = CONTENT_DIR / source_name
        if not source_path.exists():
            raise FileNotFoundError(source_path)
        pdf_path = PDF_DIR / f"{output_stem}.pdf"
        docx_path = DOCX_DIR / f"{output_stem}.docx"
        build_pdf(source_path, pdf_path, output_stem)
        build_docx(source_path, docx_path, output_stem)
        print(f"Built {output_stem}.pdf and {output_stem}.docx")


if __name__ == "__main__":
    main()
